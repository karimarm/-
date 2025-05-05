#!/usr/bin/env python
# -*- coding: utf-8 -*-

import os
from docx import Document
from PyPDF2 import PdfReader

def read_docx(file_path):
    """
    Чтение текста из файла DOCX
    
    Args:
        file_path (str): Путь к файлу
        
    Returns:
        str: Текст из файла
    """
    try:
        doc = Document(file_path)
        full_text = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text)
        
        return '\n'.join(full_text)
    except Exception as e:
        raise Exception(f"Ошибка при чтении файла DOCX: {str(e)}")

def read_pdf(file_path):
    """
    Чтение текста из файла PDF
    
    Args:
        file_path (str): Путь к файлу
        
    Returns:
        str: Текст из файла
    """
    try:
        reader = PdfReader(file_path)
        full_text = []
        
        for page in reader.pages:
            text = page.extract_text()
            if text.strip():
                full_text.append(text)
        
        return '\n'.join(full_text)
    except Exception as e:
        raise Exception(f"Ошибка при чтении файла PDF: {str(e)}")

def read_txt(file_path):
    """
    Чтение текста из текстового файла
    
    Args:
        file_path (str): Путь к файлу
        
    Returns:
        str: Текст из файла
    """
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    except UnicodeDecodeError:
        # Если не удалось прочитать в utf-8, пробуем другие кодировки
        encodings = ['windows-1251', 'latin-1', 'cp1252']
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    return file.read()
            except UnicodeDecodeError:
                continue
        raise Exception(f"Не удалось определить кодировку файла {file_path}")
    except Exception as e:
        raise Exception(f"Ошибка при чтении текстового файла: {str(e)}")

def save_to_txt(content, file_path):
    """
    Сохранение текста в текстовый файл
    
    Args:
        content (str): Содержимое для сохранения
        file_path (str): Путь к файлу
    """
    try:
        with open(file_path, 'w', encoding='utf-8') as file:
            file.write(content)
    except Exception as e:
        raise Exception(f"Ошибка при сохранении в текстовый файл: {str(e)}")

def save_to_docx(content, file_path):
    """
    Сохранение текста в файл DOCX
    
    Args:
        content (str): Содержимое для сохранения
        file_path (str): Путь к файлу
    """
    try:
        doc = Document()
        
        # Разделение текста по строкам и добавление каждой строки в отдельный параграф
        for line in content.split('\n'):
            if line.strip():
                doc.add_paragraph(line)
        
        doc.save(file_path)
    except Exception as e:
        raise Exception(f"Ошибка при сохранении в файл DOCX: {str(e)}")

def read_file(file_path):
    """
    Чтение текста из файла в зависимости от его формата
    
    Args:
        file_path (str): Путь к файлу
        
    Returns:
        str: Текст из файла
    """
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()
    
    if ext == '.docx':
        return read_docx(file_path)
    elif ext == '.pdf':
        return read_pdf(file_path)
    elif ext == '.txt':
        return read_txt(file_path)
    else:
        raise Exception(f"Неподдерживаемый формат файла: {ext}")

def save_file(content, file_path):
    """
    Сохранение текста в файл в зависимости от его формата
    
    Args:
        content (str): Содержимое для сохранения
        file_path (str): Путь к файлу
    """
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()
    
    if ext == '.docx':
        save_to_docx(content, file_path)
    elif ext == '.txt':
        save_to_txt(content, file_path)
    else:
        raise Exception(f"Неподдерживаемый формат файла для сохранения: {ext}") 